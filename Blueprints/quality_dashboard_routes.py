"""
Routes pour le dashboard qualité
Fournit les endpoints d'API pour les graphiques et analyses
"""

from flask import Blueprint, jsonify, request, render_template
from datetime import datetime, timedelta
from configuration.cosmos_config import get_call_history_container
import statistics

quality_bp = Blueprint('quality', __name__, url_prefix='/api/quality')

# Blueprint séparé pour les pages HTML (sans préfixe /api)
quality_pages_bp = Blueprint('quality_pages', __name__, url_prefix='/quality')

@quality_pages_bp.route('/report')
def quality_report_page():
    """Affiche la page de génération de rapport"""
    return render_template('quality_report.html')


def get_period_filter(period_days):
    """Génère la date de début pour filtrer les conversations"""
    end_date = datetime.utcnow()
    start_date = end_date - timedelta(days=period_days)
    return start_date


def query_conversations_by_period(period_days):
    """
    Récupère toutes les conversations terminées sur la période donnée
    """
    container = get_call_history_container()
    start_date = get_period_filter(period_days)
    start_timestamp = int(start_date.timestamp())
    
    query = """
    SELECT c.call_id, c.agent_id, c.model, c.model_family, c.status,
           c.duration_minutes, c.cost, c.tokens,
           c.user_sentiment_analysis, c.assistant_sentiment_analysis,
           c.start_time, c.end_time, c._ts
    FROM c
    WHERE c.status = 'completed'
          AND c._ts >= @start_timestamp
    ORDER BY c._ts DESC
    """
    
    parameters = [{"name": "@start_timestamp", "value": start_timestamp}]
    
    items = list(container.query_items(
        query=query,
        parameters=parameters,
        enable_cross_partition_query=True
    ))
    
    return items


@quality_bp.route('/by-family', methods=['GET'])
def get_by_family():
    """
    Retourne les sentiments cumulés groupés par famille d'agent
    2 barres par famille : User et Assistant, chacune avec sentiments empilés
    """
    period = int(request.args.get('period', 30))
    conversations = query_conversations_by_period(period)
    
    # Grouper par famille
    families = {}
    family_mapping = {}  # Pour mapper les vraies familles vers F1, F2, F3
    family_counter = 1
    
    for conv in conversations:
        # Utiliser model_family s'il existe et n'est pas vide, sinon utiliser model
        original_family = conv.get('model_family') or conv.get('model') or 'Unknown'
        if not original_family or original_family.strip() == '':
            original_family = conv.get('model') or 'Unknown'
        
        # Mapper vers F1, F2, F3, etc.
        if original_family not in family_mapping:
            family_mapping[original_family] = f'F{family_counter}'
            family_counter += 1
        
        family = family_mapping[original_family]
            
        if family not in families:
            families[family] = {
                'user_positive': [],
                'user_neutral': [],
                'user_negative': [],
                'agent_positive': [],
                'agent_neutral': [],
                'agent_negative': []
            }
        
        user_sent = conv.get('user_sentiment_analysis', {})
        agent_sent = conv.get('assistant_sentiment_analysis', {})
        
        if isinstance(user_sent, dict):
            families[family]['user_positive'].append(user_sent.get('positive', 0) or 0)
            families[family]['user_neutral'].append(user_sent.get('neutral', 0) or 0)
            families[family]['user_negative'].append(user_sent.get('negative', 0) or 0)
        
        if isinstance(agent_sent, dict):
            families[family]['agent_positive'].append(agent_sent.get('positive', 0) or 0)
            families[family]['agent_neutral'].append(agent_sent.get('neutral', 0) or 0)
            families[family]['agent_negative'].append(agent_sent.get('negative', 0) or 0)
    
    # Calculer moyennes par famille (multiplier par 100 pour avoir des pourcentages)
    labels = list(families.keys())
    user_pos = [round(statistics.mean(families[f]['user_positive']) * 100, 2) if families[f]['user_positive'] else 0 for f in labels]
    user_neut = [round(statistics.mean(families[f]['user_neutral']) * 100, 2) if families[f]['user_neutral'] else 0 for f in labels]
    user_neg = [round(statistics.mean(families[f]['user_negative']) * 100, 2) if families[f]['user_negative'] else 0 for f in labels]
    
    agent_pos = [round(statistics.mean(families[f]['agent_positive']) * 100, 2) if families[f]['agent_positive'] else 0 for f in labels]
    agent_neut = [round(statistics.mean(families[f]['agent_neutral']) * 100, 2) if families[f]['agent_neutral'] else 0 for f in labels]
    agent_neg = [round(statistics.mean(families[f]['agent_negative']) * 100, 2) if families[f]['agent_negative'] else 0 for f in labels]
    
    # Compter nombre de conversations par famille
    conv_counts = [len(families[f]['user_positive']) for f in labels]
    
    return jsonify({
        'labels': labels,
        'user_positive': user_pos,
        'user_neutral': user_neut,
        'user_negative': user_neg,
        'agent_positive': agent_pos,
        'agent_neutral': agent_neut,
        'agent_negative': agent_neg,
        'conversation_counts': conv_counts
    })


@quality_bp.route('/by-day', methods=['GET'])
def get_by_day():
    """
    Retourne les sentiments cumulés groupés par jour
    2 barres par jour : User et Assistant, chacune avec sentiments empilés
    """
    period = int(request.args.get('period', 30))
    conversations = query_conversations_by_period(period)
    
    # Grouper par jour
    days = {}
    for conv in conversations:
        timestamp = conv.get('_ts', 0)
        if timestamp:
            date = datetime.utcfromtimestamp(timestamp).strftime('%Y-%m-%d')
        else:
            continue
        
        if date not in days:
            days[date] = {
                'user_positive': [],
                'user_neutral': [],
                'user_negative': [],
                'agent_positive': [],
                'agent_neutral': [],
                'agent_negative': []
            }
        
        user_sent = conv.get('user_sentiment_analysis', {})
        agent_sent = conv.get('assistant_sentiment_analysis', {})
        
        if isinstance(user_sent, dict):
            days[date]['user_positive'].append(user_sent.get('positive', 0) or 0)
            days[date]['user_neutral'].append(user_sent.get('neutral', 0) or 0)
            days[date]['user_negative'].append(user_sent.get('negative', 0) or 0)
        
        if isinstance(agent_sent, dict):
            days[date]['agent_positive'].append(agent_sent.get('positive', 0) or 0)
            days[date]['agent_neutral'].append(agent_sent.get('neutral', 0) or 0)
            days[date]['agent_negative'].append(agent_sent.get('negative', 0) or 0)
    
    # Trier par date
    sorted_dates = sorted(days.keys())
    
    # Calculer moyennes par jour (multiplier par 100 pour avoir des pourcentages)
    user_pos = [round(statistics.mean(days[d]['user_positive']) * 100, 2) if days[d]['user_positive'] else 0 for d in sorted_dates]
    user_neut = [round(statistics.mean(days[d]['user_neutral']) * 100, 2) if days[d]['user_neutral'] else 0 for d in sorted_dates]
    user_neg = [round(statistics.mean(days[d]['user_negative']) * 100, 2) if days[d]['user_negative'] else 0 for d in sorted_dates]
    
    agent_pos = [round(statistics.mean(days[d]['agent_positive']) * 100, 2) if days[d]['agent_positive'] else 0 for d in sorted_dates]
    agent_neut = [round(statistics.mean(days[d]['agent_neutral']) * 100, 2) if days[d]['agent_neutral'] else 0 for d in sorted_dates]
    agent_neg = [round(statistics.mean(days[d]['agent_negative']) * 100, 2) if days[d]['agent_negative'] else 0 for d in sorted_dates]
    
    return jsonify({
        'labels': sorted_dates,
        'user_positive': user_pos,
        'user_neutral': user_neut,
        'user_negative': user_neg,
        'agent_positive': agent_pos,
        'agent_neutral': agent_neut,
        'agent_negative': agent_neg
    })


@quality_bp.route('/by-model', methods=['GET'])
def get_by_model():
    """
    Retourne les sentiments cumulés groupés par modèle
    2 barres par modèle : User et Assistant, chacune avec sentiments empilés
    """
    period = int(request.args.get('period', 30))
    conversations = query_conversations_by_period(period)
    
    # Grouper par modèle
    models = {}
    for conv in conversations:
        model = conv.get('model', 'Unknown')
        if model not in models:
            models[model] = {
                'user_positive': [],
                'user_neutral': [],
                'user_negative': [],
                'agent_positive': [],
                'agent_neutral': [],
                'agent_negative': []
            }
        
        user_sent = conv.get('user_sentiment_analysis', {})
        agent_sent = conv.get('assistant_sentiment_analysis', {})
        
        if isinstance(user_sent, dict):
            models[model]['user_positive'].append(user_sent.get('positive', 0) or 0)
            models[model]['user_neutral'].append(user_sent.get('neutral', 0) or 0)
            models[model]['user_negative'].append(user_sent.get('negative', 0) or 0)
        
        if isinstance(agent_sent, dict):
            models[model]['agent_positive'].append(agent_sent.get('positive', 0) or 0)
            models[model]['agent_neutral'].append(agent_sent.get('neutral', 0) or 0)
            models[model]['agent_negative'].append(agent_sent.get('negative', 0) or 0)
    
    # Calculer moyennes par modèle (multiplier par 100 pour avoir des pourcentages)
    labels = list(models.keys())
    user_pos = [round(statistics.mean(models[m]['user_positive']) * 100, 2) if models[m]['user_positive'] else 0 for m in labels]
    user_neut = [round(statistics.mean(models[m]['user_neutral']) * 100, 2) if models[m]['user_neutral'] else 0 for m in labels]
    user_neg = [round(statistics.mean(models[m]['user_negative']) * 100, 2) if models[m]['user_negative'] else 0 for m in labels]
    
    agent_pos = [round(statistics.mean(models[m]['agent_positive']) * 100, 2) if models[m]['agent_positive'] else 0 for m in labels]
    agent_neut = [round(statistics.mean(models[m]['agent_neutral']) * 100, 2) if models[m]['agent_neutral'] else 0 for m in labels]
    agent_neg = [round(statistics.mean(models[m]['agent_negative']) * 100, 2) if models[m]['agent_negative'] else 0 for m in labels]
    
    return jsonify({
        'labels': labels,
        'user_positive': user_pos,
        'user_neutral': user_neut,
        'user_negative': user_neg,
        'agent_positive': agent_pos,
        'agent_neutral': agent_neut,
        'agent_negative': agent_neg
    })


@quality_bp.route('/sentiment-trend', methods=['GET'])
def get_sentiment_trend():
    """
    Retourne l'évolution des sentiments au fil du temps avec indices de tendance
    """
    period = int(request.args.get('period', 30))
    conversations = query_conversations_by_period(period)
    
    # Grouper par jour et calculer moyennes de sentiments
    days = {}
    for conv in conversations:
        timestamp = conv.get('_ts', 0)
        if timestamp:
            date = datetime.utcfromtimestamp(timestamp).strftime('%Y-%m-%d')
        else:
            continue
        
        if date not in days:
            days[date] = {
                'user_positive': [],
                'user_neutral': [],
                'user_negative': [],
                'agent_positive': [],
                'agent_neutral': [],
                'agent_negative': []
            }
        
        # Séparer user et assistant sentiment
        user_sent = conv.get('user_sentiment_analysis', {})
        agent_sent = conv.get('assistant_sentiment_analysis', {})
        
        if isinstance(user_sent, dict):
            days[date]['user_positive'].append((user_sent.get('positive', 0) or 0) * 100)
            days[date]['user_neutral'].append((user_sent.get('neutral', 0) or 0) * 100)
            days[date]['user_negative'].append((user_sent.get('negative', 0) or 0) * 100)
        
        if isinstance(agent_sent, dict):
            days[date]['agent_positive'].append((agent_sent.get('positive', 0) or 0) * 100)
            days[date]['agent_neutral'].append((agent_sent.get('neutral', 0) or 0) * 100)
            days[date]['agent_negative'].append((agent_sent.get('negative', 0) or 0) * 100)
    
    # Calculer moyennes par jour
    sorted_dates = sorted(days.keys())
    labels = sorted_dates
    user_positive = [statistics.mean(days[d]['user_positive']) if days[d]['user_positive'] else 0 for d in sorted_dates]
    user_neutral = [statistics.mean(days[d]['user_neutral']) if days[d]['user_neutral'] else 0 for d in sorted_dates]
    user_negative = [statistics.mean(days[d]['user_negative']) if days[d]['user_negative'] else 0 for d in sorted_dates]
    agent_positive = [statistics.mean(days[d]['agent_positive']) if days[d]['agent_positive'] else 0 for d in sorted_dates]
    agent_neutral = [statistics.mean(days[d]['agent_neutral']) if days[d]['agent_neutral'] else 0 for d in sorted_dates]
    agent_negative = [statistics.mean(days[d]['agent_negative']) if days[d]['agent_negative'] else 0 for d in sorted_dates]
    
    return jsonify({
        'labels': labels,
        'user_positive': user_positive,
        'user_neutral': user_neutral,
        'user_negative': user_negative,
        'agent_positive': agent_positive,
        'agent_neutral': agent_neutral,
        'agent_negative': agent_negative
    })


@quality_bp.route('/correlations', methods=['GET'])
def get_correlations():
    """
    Calcule les corrélations entre sentiments agent et user
    """
    period = int(request.args.get('period', 30))
    conversations = query_conversations_by_period(period)
    
    # Collecter les données pour scatter plots
    pos_agent = []
    neut_agent = []
    neg_agent = []
    pos_user = []
    
    for conv in conversations:
        user_sent = conv.get('user_sentiment_analysis', {})
        agent_sent = conv.get('assistant_sentiment_analysis', {})
        
        if isinstance(user_sent, dict) and isinstance(agent_sent, dict):
            pos_u = (user_sent.get('positive', 0) or 0) * 100
            pos_a = (agent_sent.get('positive', 0) or 0) * 100
            neut_a = (agent_sent.get('neutral', 0) or 0) * 100
            neg_a = (agent_sent.get('negative', 0) or 0) * 100
            
            if pos_u > 0 or pos_a > 0 or neut_a > 0 or neg_a > 0:
                pos_agent.append(pos_a)
                neut_agent.append(neut_a)
                neg_agent.append(neg_a)
                pos_user.append(pos_u)
    
    # Calculer corrélations
    def pearson_correlation(x, y):
        if len(x) != len(y) or len(x) < 2:
            return 0
        
        n = len(x)
        sum_x = sum(x)
        sum_y = sum(y)
        sum_xy = sum(xi * yi for xi, yi in zip(x, y))
        sum_x2 = sum(xi * xi for xi in x)
        sum_y2 = sum(yi * yi for yi in y)
        
        numerator = n * sum_xy - sum_x * sum_y
        denominator = ((n * sum_x2 - sum_x ** 2) * (n * sum_y2 - sum_y ** 2)) ** 0.5
        
        return numerator / denominator if denominator != 0 else 0
    
    corr_pos_pos = pearson_correlation(pos_agent, pos_user)
    corr_neut_pos = pearson_correlation(neut_agent, pos_user)
    corr_neg_pos = pearson_correlation(neg_agent, pos_user)
    
    # Préparer scatter data
    scatter_pos_pos = [{'x': pos_agent[i], 'y': pos_user[i]} for i in range(len(pos_agent))]
    scatter_neut_pos = [{'x': neut_agent[i], 'y': pos_user[i]} for i in range(len(neut_agent))]
    scatter_neg_pos = [{'x': neg_agent[i], 'y': pos_user[i]} for i in range(len(neg_agent))]
    
    return jsonify({
        'correlations': {
            'pos_pos': corr_pos_pos,
            'neut_pos': corr_neut_pos,
            'neg_pos': corr_neg_pos
        },
        'scatter_data': {
            'pos_pos': scatter_pos_pos,
            'neut_pos': scatter_neut_pos,
            'neg_pos': scatter_neg_pos
        }
    })


@quality_bp.route('/generate-report', methods=['GET'])
def generate_report():
    """
    Génère un rapport d'analyse détaillé avec GPT-4o-mini
    """
    try:
        import os
        import sys
        print(f"[DEBUG] Python executable: {sys.executable}")
        print(f"[DEBUG] Python version: {sys.version}")
        print("[DEBUG] Attempting to import openai...")
        from openai import AzureOpenAI
        print("[DEBUG] ✅ openai imported successfully")
        
        period = int(request.args.get('period', 30))
        
        # Collecter toutes les données du dashboard
        conversations = query_conversations_by_period(period)
        
        # Statistiques générales
        total_conversations = len(conversations)
        
        # Calcul de moyennes globales
        all_user_pos = []
        all_user_neut = []
        all_user_neg = []
        all_agent_pos = []
        all_agent_neut = []
        all_agent_neg = []
        
        for conv in conversations:
            user_sent = conv.get('user_sentiment_analysis', {})
            agent_sent = conv.get('assistant_sentiment_analysis', {})
            
            if isinstance(user_sent, dict):
                all_user_pos.append((user_sent.get('positive', 0) or 0) * 100)
                all_user_neut.append((user_sent.get('neutral', 0) or 0) * 100)
                all_user_neg.append((user_sent.get('negative', 0) or 0) * 100)
            
            if isinstance(agent_sent, dict):
                all_agent_pos.append((agent_sent.get('positive', 0) or 0) * 100)
                all_agent_neut.append((agent_sent.get('neutral', 0) or 0) * 100)
                all_agent_neg.append((agent_sent.get('negative', 0) or 0) * 100)
        
        # Statistiques par famille
        families = {}
        family_mapping = {}
        family_counter = 1
        
        for conv in conversations:
            original_family = conv.get('model_family') or conv.get('model') or 'Unknown'
            if not original_family or original_family.strip() == '':
                original_family = conv.get('model') or 'Unknown'
            
            if original_family not in family_mapping:
                family_mapping[original_family] = f'F{family_counter}'
                family_counter += 1
            
            family = family_mapping[original_family]
            
            if family not in families:
                families[family] = {
                    'count': 0,
                    'user_pos': [], 'user_neut': [], 'user_neg': [],
                    'agent_pos': [], 'agent_neut': [], 'agent_neg': []
                }
            
            families[family]['count'] += 1
            
            user_sent = conv.get('user_sentiment_analysis', {})
            agent_sent = conv.get('assistant_sentiment_analysis', {})
            
            if isinstance(user_sent, dict):
                families[family]['user_pos'].append((user_sent.get('positive', 0) or 0) * 100)
                families[family]['user_neut'].append((user_sent.get('neutral', 0) or 0) * 100)
                families[family]['user_neg'].append((user_sent.get('negative', 0) or 0) * 100)
            
            if isinstance(agent_sent, dict):
                families[family]['agent_pos'].append((agent_sent.get('positive', 0) or 0) * 100)
                families[family]['agent_neut'].append((agent_sent.get('neutral', 0) or 0) * 100)
                families[family]['agent_neg'].append((agent_sent.get('negative', 0) or 0) * 100)
        
        # Statistiques par jour
        days = {}
        for conv in conversations:
            timestamp = conv.get('_ts', 0)
            if timestamp:
                date = datetime.utcfromtimestamp(timestamp).strftime('%Y-%m-%d')
                if date not in days:
                    days[date] = {
                        'count': 0,
                        'user_pos': [], 'user_neut': [], 'user_neg': [],
                        'agent_pos': [], 'agent_neut': [], 'agent_neg': []
                    }
                
                days[date]['count'] += 1
                
                user_sent = conv.get('user_sentiment_analysis', {})
                agent_sent = conv.get('assistant_sentiment_analysis', {})
                
                if isinstance(user_sent, dict):
                    days[date]['user_pos'].append((user_sent.get('positive', 0) or 0) * 100)
                    days[date]['user_neut'].append((user_sent.get('neutral', 0) or 0) * 100)
                    days[date]['user_neg'].append((user_sent.get('negative', 0) or 0) * 100)
                
                if isinstance(agent_sent, dict):
                    days[date]['agent_pos'].append((agent_sent.get('positive', 0) or 0) * 100)
                    days[date]['agent_neut'].append((agent_sent.get('neutral', 0) or 0) * 100)
                    days[date]['agent_neg'].append((agent_sent.get('negative', 0) or 0) * 100)
        
        # Préparer le résumé des données pour GPT
        family_stats = []
        for family, data in families.items():
            family_stats.append({
                'family': family,
                'conversations': data['count'],
                'user_positive_avg': round(statistics.mean(data['user_pos']), 2) if data['user_pos'] else 0,
                'user_neutral_avg': round(statistics.mean(data['user_neut']), 2) if data['user_neut'] else 0,
                'user_negative_avg': round(statistics.mean(data['user_neg']), 2) if data['user_neg'] else 0,
                'agent_positive_avg': round(statistics.mean(data['agent_pos']), 2) if data['agent_pos'] else 0,
                'agent_neutral_avg': round(statistics.mean(data['agent_neut']), 2) if data['agent_neut'] else 0,
                'agent_negative_avg': round(statistics.mean(data['agent_neg']), 2) if data['agent_neg'] else 0
            })
        
        day_stats = []
        for date, data in sorted(days.items()):
            day_stats.append({
                'date': date,
                'conversations': data['count'],
                'user_positive_avg': round(statistics.mean(data['user_pos']), 2) if data['user_pos'] else 0,
                'user_negative_avg': round(statistics.mean(data['user_neg']), 2) if data['user_neg'] else 0,
                'agent_positive_avg': round(statistics.mean(data['agent_pos']), 2) if data['agent_pos'] else 0,
                'agent_negative_avg': round(statistics.mean(data['agent_neg']), 2) if data['agent_neg'] else 0
            })
        
        # Construire le prompt pour GPT-4o-mini
        prompt = f"""Tu es un analyste de données spécialisé dans l'analyse de sentiments conversationnels.

Voici les données d'analyse de sentiment sur les {period} derniers jours :

STATISTIQUES GLOBALES :
- Nombre total de conversations : {total_conversations}
- Sentiment utilisateur moyen : Positif {round(statistics.mean(all_user_pos), 2) if all_user_pos else 0}%, Neutre {round(statistics.mean(all_user_neut), 2) if all_user_neut else 0}%, Négatif {round(statistics.mean(all_user_neg), 2) if all_user_neg else 0}%
- Sentiment assistant moyen : Positif {round(statistics.mean(all_agent_pos), 2) if all_agent_pos else 0}%, Neutre {round(statistics.mean(all_agent_neut), 2) if all_agent_neut else 0}%, Négatif {round(statistics.mean(all_agent_neg), 2) if all_agent_neg else 0}%

ANALYSE PAR FAMILLE D'AGENTS :
{chr(10).join([f"- {f['family']} : {f['conversations']} conversations | User: +{f['user_positive_avg']}% ={f['user_neutral_avg']}% -{f['user_negative_avg']}% | Agent: +{f['agent_positive_avg']}% ={f['agent_neutral_avg']}% -{f['agent_negative_avg']}%" for f in family_stats])}

ÉVOLUTION QUOTIDIENNE (derniers 7 jours) :
{chr(10).join([f"- {d['date']} : {d['conversations']} conv | User: +{d['user_positive_avg']}% -{d['user_negative_avg']}% | Agent: +{d['agent_positive_avg']}% -{d['agent_negative_avg']}%" for d in day_stats[-7:]])}

INSTRUCTIONS :
Génère un rapport d'analyse structuré en Markdown avec les sections suivantes :

## Analyse des Sentiments - Rapport sur {period} jours

### 1. Vue d'Ensemble
Présente les statistiques globales de manière factuelle et DÉTAILLÉE (minimum 5000 caractères). 
- Décris chaque métrique en profondeur
- Analyse les distributions statistiques
- Compare avec les moyennes attendues
- Explique les écarts-types et variances
- Inclus des analyses de corrélation entre métriques
- Suggère des graphiques : histogramme des sentiments globaux, camembert de répartition positif/neutre/négatif

### 2. Analyse par Famille d'Agents
Compare EXHAUSTIVEMENT les performances des différentes familles (minimum 5000 caractères).
- Analyse détaillée de chaque famille individuellement
- Comparaisons croisées entre toutes les familles
- Identifie les différences significatives (>10 points de pourcentage)
- Analyse les tendances par famille
- Calcule les écarts-types par famille
- Suggère des graphiques : graphiques en barres groupées par famille, courbes de tendance par famille

### 3. Évolution Temporelle
Analyse APPROFONDIE des tendances quotidiennes (minimum 5000 caractères).
- Détaille jour par jour les évolutions
- Identifie les événements marquants avec contexte détaillé :
  * Fortes hausses de sentiment positif (>15% d'augmentation)
  * Fortes baisses (>15% de diminution)
  * Pics d'activité (jours avec >20% du volume total)
- Analyse les patterns hebdomadaires
- Calcule les moyennes mobiles sur 3 et 7 jours
- Identifie les cycles et saisonnalités
- Suggère des graphiques : courbes temporelles multi-axes, heatmap calendrier des sentiments

### 4. Points d'Attention
Liste DÉTAILLÉE et EXHAUSTIVE des métriques (minimum 5000 caractères).
- Analyse statistique complète (écart-type >1.5)
- Identification des outliers avec contexte
- Analyse des anomalies temporelles
- Corrélations entre différentes métriques
- Patterns inhabituels détectés
- Suggère des graphiques : box plots des distributions, scatter plots des corrélations

### 5. Conclusion Générale
Résume COMPLÈTEMENT tous les constats (minimum 3000 caractères).
- Synthèse détaillée de chaque section
- Liens entre les différentes analyses
- Constats principaux en 10-15 points factuels
- Vue d'ensemble des tendances globales
- Suggère un graphique récapitulatif : dashboard synthétique multi-métriques

IMPORTANT :
- CHAQUE SECTION DOIT CONTENIR AU MINIMUM 5000 CARACTÈRES (sauf conclusion: 3000)
- Sois TRÈS DÉTAILLÉ et EXHAUSTIF dans chaque analyse
- Développe chaque point avec des explications complètes
- Utilise des paragraphes longs et détaillés
- Suggère des graphiques pertinents pour CHAQUE section
- Mets en gras les chiffres clés avec **
- Reste factuel mais TRÈS COMPLET dans l'analyse des données
"""
        
        # Appeler GPT-4o-mini
        client = AzureOpenAI(
            api_key=os.environ.get("AZURE_OPENAI_SUMMARY_KEY"),
            api_version=os.environ.get("AZURE_OPENAI_SUMMARY_API_VERSION", "2024-02-15-preview"),
            azure_endpoint=os.environ.get("AZURE_OPENAI_SUMMARY_ENDPOINT")
        )
        
        deployment_name = os.environ.get("AZURE_OPENAI_GPT5_MINI_DEPLOYMENT", "gpt-4o-mini")
        
        print(f"[DEBUG] Calling Azure OpenAI with deployment: {deployment_name}")
        print(f"[DEBUG] Prompt length: {len(prompt)} characters")
        
        response = client.chat.completions.create(
            model=deployment_name,
            messages=[
                {"role": "system", "content": "Tu es un analyste de données expert en analyse statistique de sentiments."},
                {"role": "user", "content": prompt}
            ],
            max_completion_tokens=32000
        )
        
        print(f"[DEBUG] Response received: {response}")
        print(f"[DEBUG] Choices: {response.choices}")
        
        report = response.choices[0].message.content
        
        print(f"[DEBUG] Report content length: {len(report) if report else 0}")
        print(f"[DEBUG] Report preview: {report[:200] if report else 'EMPTY'}")
        
        return jsonify({
            'success': True,
            'report': report
        })
        
    except Exception as e:
        print(f"Erreur génération rapport: {str(e)}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@quality_bp.route('/export-report', methods=['POST'])
def export_report():
    """
    Exporte un rapport en format Word ou PDF
    """
    try:
        data = request.json
        report_markdown = data.get('report', '')
        format_type = data.get('format', 'docx')
        period_days = data.get('period', 30)  # Récupérer la période
        format_type = data.get('format', 'docx')
        
        if format_type == 'docx':
            from docx import Document
            from docx.shared import Pt, RGBColor, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            from io import BytesIO
            from datetime import datetime
            import os
            
            # Créer un document Word
            doc = Document()
            
            # ==================== PAGE DE GARDE ====================
            
            # Ajouter le logo en haut (si le fichier existe)
            logo_path = os.path.join('static', 'images', 'logo.png')
            if os.path.exists(logo_path):
                logo_para = doc.add_paragraph()
                logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = logo_para.add_run()
                run.add_picture(logo_path, width=Inches(2.5))
            
            # Espacement après le logo
            doc.add_paragraph()
            
            # Bandeau violet avec titre principal
            # Créer un paragraphe avec fond violet
            title = doc.add_heading('Waka Burkina Faso', level=0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Ajouter un fond violet au paragraphe
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), '667eea')  # Violet Waka
            title._element.get_or_add_pPr().append(shading_elm)
            
            for run in title.runs:
                run.font.size = Pt(36)
                run.font.color.rgb = RGBColor(255, 255, 255)  # Blanc sur fond violet
                run.font.bold = True
            
            # Espacement
            doc.add_paragraph()
            
            # Sous-titre
            subtitle = doc.add_heading('Rapport d\'Analyse Qualité', level=1)
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in subtitle.runs:
                run.font.size = Pt(24)
                run.font.color.rgb = RGBColor(255, 165, 0)  # Orange
            
            # Période d'analyse
            period_para = doc.add_paragraph(f'Période analysée: {period_days} jours')
            period_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in period_para.runs:
                run.font.size = Pt(16)
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 215, 0)  # Jaune/Or
            
            # Espacement
            doc.add_paragraph()
            
            # Centre de contact
            center = doc.add_paragraph('Centre de Contact')
            center.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in center.runs:
                run.font.size = Pt(18)
                run.font.bold = True
                run.font.color.rgb = RGBColor(102, 126, 234)  # Violet
            
            # Espacement
            doc.add_paragraph()
            doc.add_paragraph()
            
            # Date
            date_para = doc.add_paragraph(f'Date: {datetime.now().strftime("%d/%m/%Y")}')
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in date_para.runs:
                run.font.size = Pt(14)
            
            # Espacement
            doc.add_paragraph()
            
            # Rédigé par
            author = doc.add_paragraph('Rédigé par')
            author.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in author.runs:
                run.font.size = Pt(14)
                run.font.italic = True
            
            team = doc.add_paragraph('Équipe Qualité')
            team.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in team.runs:
                run.font.size = Pt(16)
                run.font.bold = True
                run.font.color.rgb = RGBColor(102, 126, 234)  # Violet Waka
            
            team2 = doc.add_paragraph('Centre de Contact')
            team2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in team2.runs:
                run.font.size = Pt(16)
                run.font.bold = True
                run.font.color.rgb = RGBColor(102, 126, 234)  # Violet Waka
            
            # Saut de page après la page de garde
            doc.add_page_break()
            
            # ==================== CONTENU DU RAPPORT ====================
            # Parser le markdown et créer le document
            lines = report_markdown.split('\n')
            for line in lines:
                if line.startswith('## '):
                    # Titre niveau 2 avec couleur Waka
                    heading = doc.add_heading(line[3:], level=1)
                    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for run in heading.runs:
                        run.font.color.rgb = RGBColor(102, 126, 234)  # Violet Waka
                elif line.startswith('### '):
                    # Titre niveau 3 avec couleur Waka
                    heading = doc.add_heading(line[4:], level=2)
                    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for run in heading.runs:
                        run.font.color.rgb = RGBColor(74, 144, 226)  # Bleu Waka
                elif line.strip() and not line.startswith('#'):
                    # Paragraphe normal
                    p = doc.add_paragraph()
                    
                    # Gérer le texte en gras **texte**
                    parts = line.split('**')
                    for i, part in enumerate(parts):
                        if i % 2 == 0:
                            # Texte normal
                            p.add_run(part)
                        else:
                            # Texte en gras avec couleur Waka
                            run = p.add_run(part)
                            run.bold = True
                            run.font.color.rgb = RGBColor(102, 126, 234)  # Violet Waka
            
            # Sauvegarder dans un buffer
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            from flask import send_file
            return send_file(
                buffer,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                as_attachment=True,
                download_name=f'rapport_qualite.docx'
            )
        
        elif format_type == 'pdf':
            from reportlab.lib.pagesizes import A4
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, PageBreak, Table, TableStyle
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.lib.colors import HexColor, white
            from reportlab.lib.enums import TA_CENTER
            from reportlab.lib import colors
            from io import BytesIO
            from datetime import datetime
            import os
            
            buffer = BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=A4)
            styles = getSampleStyleSheet()
            story = []
            
            # ==================== PAGE DE GARDE PDF ====================
            
            # Logo (si existe)
            logo_path = os.path.join('static', 'images', 'logo.png')
            if os.path.exists(logo_path):
                logo = Image(logo_path, width=2.5*inch, height=2.5*inch)
                story.append(logo)
                story.append(Spacer(1, 0.5*inch))
            
            # Style pour le titre principal avec fond violet
            cover_title_style = ParagraphStyle(
                'CoverTitle',
                parent=styles['Title'],
                fontSize=36,
                textColor=white,
                alignment=TA_CENTER,
                spaceAfter=20,
                leading=42
            )
            
            # Créer un tableau pour simuler le bandeau violet
            title_data = [['Waka Burkina Faso']]
            title_table = Table(title_data, colWidths=[6*inch])
            title_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, -1), HexColor('#667eea')),
                ('TEXTCOLOR', (0, 0), (-1, -1), white),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, -1), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 36),
                ('TOPPADDING', (0, 0), (-1, -1), 20),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 20),
            ]))
            story.append(title_table)
            story.append(Spacer(1, 0.5*inch))
            
            # Sous-titre
            subtitle_style = ParagraphStyle(
                'CoverSubtitle',
                parent=styles['Heading1'],
                fontSize=24,
                textColor=HexColor('#FFA500'),  # Orange
                alignment=TA_CENTER,
                spaceAfter=20
            )
            story.append(Paragraph("Rapport d'Analyse Qualité", subtitle_style))
            story.append(Spacer(1, 0.3*inch))
            
            # Période
            period_style = ParagraphStyle(
                'CoverPeriod',
                parent=styles['Normal'],
                fontSize=16,
                textColor=HexColor('#FFD700'),  # Jaune/Or
                alignment=TA_CENTER,
                spaceAfter=12,
                fontName='Helvetica-Bold'
            )
            story.append(Paragraph(f"Période analysée: {period_days} jours", period_style))
            story.append(Spacer(1, 0.3*inch))
            
            # Centre de contact
            center_style = ParagraphStyle(
                'CoverCenter',
                parent=styles['Normal'],
                fontSize=18,
                textColor=HexColor('#667eea'),  # Violet
                alignment=TA_CENTER,
                spaceAfter=12,
                fontName='Helvetica-Bold'
            )
            story.append(Paragraph("Centre de Contact", center_style))
            story.append(Spacer(1, 0.5*inch))
            
            # Date
            date_style = ParagraphStyle(
                'CoverDate',
                parent=styles['Normal'],
                fontSize=14,
                alignment=TA_CENTER,
                spaceAfter=12
            )
            story.append(Paragraph(f"Date: {datetime.now().strftime('%d/%m/%Y')}", date_style))
            story.append(Spacer(1, 0.3*inch))
            
            # Rédigé par
            author_style = ParagraphStyle(
                'CoverAuthor',
                parent=styles['Normal'],
                fontSize=14,
                alignment=TA_CENTER,
                spaceAfter=8,
                fontName='Helvetica-Oblique'
            )
            story.append(Paragraph("Rédigé par", author_style))
            
            team_style = ParagraphStyle(
                'CoverTeam',
                parent=styles['Normal'],
                fontSize=16,
                textColor=HexColor('#667eea'),  # Violet
                alignment=TA_CENTER,
                spaceAfter=6,
                fontName='Helvetica-Bold'
            )
            story.append(Paragraph("Équipe Qualité", team_style))
            story.append(Paragraph("Centre de Contact", team_style))
            
            # Saut de page après la page de garde
            story.append(PageBreak())
            
            # ==================== CONTENU DU RAPPORT ====================
            
            # Créer des styles personnalisés pour le contenu
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=16,
                textColor=HexColor('#667eea'),  # Violet
                spaceAfter=12
            )
            
            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading2'],
                fontSize=14,
                textColor=HexColor('#FFA500'),  # Orange
                spaceAfter=10
            )
            
            # Parser le markdown
            lines = report_markdown.split('\n')
            for line in lines:
                if line.startswith('## '):
                    story.append(Paragraph(line[3:], title_style))
                    story.append(Spacer(1, 0.2*inch))
                elif line.startswith('### '):
                    story.append(Paragraph(line[4:], heading_style))
                    story.append(Spacer(1, 0.1*inch))
                elif line.strip() and not line.startswith('#'):
                    # Remplacer ** par <b> pour le gras
                    formatted_line = line.replace('**', '<b>', 1)
                    while '**' in formatted_line:
                        formatted_line = formatted_line.replace('**', '</b>', 1)
                        if '**' in formatted_line:
                            formatted_line = formatted_line.replace('**', '<b>', 1)
                    
                    story.append(Paragraph(formatted_line, styles['Normal']))
                    story.append(Spacer(1, 0.1*inch))
            
            doc.build(story)
            buffer.seek(0)
            
            from flask import send_file
            return send_file(
                buffer,
                mimetype='application/pdf',
                as_attachment=True,
                download_name=f'rapport_qualite.pdf'
            )
        
        else:
            return jsonify({'success': False, 'error': 'Format non supporté'}), 400
            
    except Exception as e:
        print(f"Erreur export rapport: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500
