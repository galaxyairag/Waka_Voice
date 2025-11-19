# tools/tool_cv.py
"""
Outil de génération de CV basé sur l'analyse de conversation
Lit l'historique depuis Cosmos DB et génère un CV professionnel via Azure OpenAI
"""

import os
import re
import uuid
import requests
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from azure.communication.email import EmailClient

load_dotenv()

# Configuration Azure OpenAI pour génération CV
AZURE_OPENAI_SUMMARY_ENDPOINT = os.getenv("AZURE_OPENAI_SUMMARY_ENDPOINT", "")
AZURE_OPENAI_SUMMARY_KEY = os.getenv("AZURE_OPENAI_SUMMARY_KEY", "")
AZURE_OPENAI_SUMMARY_DEPLOYMENT = os.getenv("AZURE_OPENAI_SUMMARY_DEPLOYMENT", "gpt-5-mini")
AZURE_OPENAI_SUMMARY_API_VERSION = os.getenv("AZURE_OPENAI_SUMMARY_API_VERSION", "2024-08-01-preview")

# Configuration Azure Communication Services Email
AZURE_COMMUNICATION_EMAIL_CONNECTION_STRING = os.getenv("AZURE_COMMUNICATION_EMAIL_CONNECTION_STRING", "")
AZURE_COMMUNICATION_EMAIL_SENDER = os.getenv("AZURE_COMMUNICATION_EMAIL_SENDER", "")


def get_tool_definition():
    """Retourne la définition du tool create_cv"""
    return {
        "type": "function",
        "name": "create_cv",
        "description": """Génère un CV professionnel au format Word à partir des informations du candidat collectées pendant la conversation. 
        
IMPORTANT: Tu dois TOUJOURS collecter ces informations une par une en posant des questions au candidat avant d'appeler ce tool:
- Nom complet
- Email et téléphone
- Adresse
- Objectif professionnel / poste recherché
- Expériences professionnelles (poste, entreprise, dates, responsabilités)
- Formations et diplômes
- Compétences
- Langues

Une fois toutes les informations collectées, appelle ce tool avec le call_id actuel de la session.""",
        "parameters": {
            "type": "object",
            "properties": {
                "call_id": {
                    "type": "string",
                    "description": "ID de l'appel Voice Live en cours (fourni automatiquement par le système)"
                },
                "email": {
                    "type": "string",
                    "description": "Adresse email où envoyer le CV généré"
                },
                "style": {
                    "type": "string",
                    "description": "Style visuel du CV",
                    "enum": ["classique", "moderne", "minimaliste"],
                    "default": "moderne"
                },
                "color": {
                    "type": "string",
                    "description": "Couleur principale du CV",
                    "enum": ["bleu", "vert", "gris", "rouge"],
                    "default": "bleu"
                }
            },
            "required": ["call_id", "email"]
        }
    }


def get_conversation_history_from_cosmosdb(call_id):
    """
    Récupère l'historique complet de la conversation depuis Cosmos DB.
    
    Args:
        call_id: ID de l'appel Voice Live
    
    Returns:
        list: Liste des messages formatés ou None si erreur
    """
    try:
        from configuration.cosmos_config import get_conversation_history
        
        print(f"📚 Récupération de l'historique depuis Cosmos DB pour call_id: {call_id}")
        
        # Récupérer le document de conversation
        conversation_doc = get_conversation_history(call_id)
        
        if not conversation_doc:
            print(f"⚠️ Aucun historique trouvé dans Cosmos DB pour call_id: {call_id}")
            return None
        
        # Extraire les messages
        messages = conversation_doc.get('conversation', [])
        
        if not messages:
            print(f"⚠️ Aucun message trouvé dans la conversation {call_id}")
            return None
        
        # Formater l'historique pour le prompt
        conversation_history = []
        for msg in messages:
            msg_type = msg.get('type', 'user')
            content = msg.get('content', '')
            timestamp = msg.get('timestamp', '')
            
            # Convertir type en role pour cohérence
            if msg_type == 'agent':
                role = 'assistant'
            elif msg_type == 'system':
                role = 'system'
            elif msg_type == 'tool':
                role = 'function'
            else:
                role = 'user'
            
            conversation_history.append({
                "role": role,
                "content": content,
                "timestamp": timestamp
            })
        
        print(f"✅ {len(conversation_history)} messages récupérés depuis Cosmos DB")
        return conversation_history
        
    except Exception as e:
        print(f"❌ Erreur lors de la récupération depuis Cosmos DB: {str(e)}")
        import traceback
        traceback.print_exc()
        return None


def generate_cv_from_conversation(conversation_history, style, color):
    """
    Génère un CV en Markdown à partir de l'historique de conversation via Azure OpenAI.
    
    Args:
        conversation_history: Liste des messages de la conversation
        style: Style du CV (classique, moderne, minimaliste)
        color: Couleur du CV (bleu, vert, gris, rouge)
    
    Returns:
        str: CV au format Markdown ou None si erreur
    """
    try:
        # Construire l'historique formaté pour le prompt
        conversation_text = "\n".join([
            f"[{msg['role'].upper()}]: {msg['content']}"
            for msg in conversation_history
        ])
        
        # Prompt pour l'IA
        system_message = """Tu es un expert en rédaction de CV professionnels modernes et attractifs. 
Tu analyses des conversations et tu génères des CV complets en Markdown avec une mise en page sophistiquée.
Tu NE dois générer QUE du Markdown, SANS commentaires ni explications."""
        
        user_message = f"""Voici l'historique complet d'une conversation où un candidat a fourni ses informations professionnelles:

{conversation_text}

Analyse cet échange et génère un CV professionnel MODERNE et VISUELLEMENT ATTRACTIF en Markdown avec:

📋 STYLE VISUEL: {style}
🎨 COULEUR THÉMATIQUE: {color}

📐 STRUCTURE OBLIGATOIRE (avec ces sections dans cet ordre):

# [NOM PRÉNOM EN MAJUSCULES]
### [Titre professionnel accrocheur]

---

## 📞 CONTACT
- 📧 Email: [email]
- 📱 Téléphone: [téléphone]
- 📍 Localisation: [ville, pays]

---

## 💼 PROFIL PROFESSIONNEL
[Un paragraphe percutant de 3-4 lignes résumant l'expertise, les années d'expérience, et la valeur ajoutée]

---

## 🎯 COMPÉTENCES CLÉS

### Compétences Techniques
- **[Catégorie 1]:** [Compétence 1] • [Compétence 2] • [Compétence 3]
- **[Catégorie 2]:** [Compétence 1] • [Compétence 2] • [Compétence 3]

### Compétences Transversales
- [Compétence] • [Compétence] • [Compétence]

---

## 💼 EXPÉRIENCE PROFESSIONNELLE

### [Poste] | [Entreprise]
**[Dates - Dates] • [Ville]**

- ✅ [Réalisation/Responsabilité avec résultat chiffré si possible]
- ✅ [Réalisation/Responsabilité avec impact]
- ✅ [Réalisation/Responsabilité avec technologies utilisées]

[Répéter pour chaque expérience]

---

## 🎓 FORMATION

### [Diplôme] | [Établissement]
**[Année] • [Ville]**
- [Mention/Spécialisation si applicable]

---

## 🌍 LANGUES
- **[Langue]:** [Niveau] (ex: Natif, Courant, Intermédiaire, Notions)

---

## 🏆 CERTIFICATIONS & DISTINCTIONS (si applicable)
- [Certification] - [Organisme] ([Année])

---

## 🎨 CENTRES D'INTÉRÊT (si mentionné)
[Intérêt 1] • [Intérêt 2] • [Intérêt 3]

---

⚡ CONSIGNES CRITIQUES:
- Utilise des ÉMOJIS pour chaque section (📞 💼 🎯 🎓 etc.)
- Ajoute des puces visuelles (• ✅) pour les listes
- Utilise le **gras** pour les mots-clés importants
- Ajoute des séparateurs (---) entre sections
- Quantifie les résultats quand possible (chiffres, pourcentages)
- Maximum 2 pages de contenu (dense mais lisible)
- Développe TOUT ce qui est mentionné dans la conversation

Génère UNIQUEMENT le Markdown du CV, sans aucun commentaire ni explication."""

        # Appel Azure OpenAI
        headers = {
            "Content-Type": "application/json",
            "api-key": AZURE_OPENAI_SUMMARY_KEY
        }
        
        endpoint = f"{AZURE_OPENAI_SUMMARY_ENDPOINT}/openai/deployments/{AZURE_OPENAI_SUMMARY_DEPLOYMENT}/chat/completions?api-version={AZURE_OPENAI_SUMMARY_API_VERSION}"
        
        payload = {
            "messages": [
                {"role": "system", "content": system_message},
                {"role": "user", "content": user_message}
            ],
            "max_completion_tokens": 32000
            # Note: temperature n'est pas supporté par tous les modèles (ex: gpt-5-mini)
            # Utiliser la valeur par défaut (1.0)
        }
        
        print(f"🌐 Appel Azure OpenAI: {AZURE_OPENAI_SUMMARY_DEPLOYMENT}")
        print(f"🔗 Endpoint: {endpoint}")
        response = requests.post(endpoint, json=payload, headers=headers, timeout=120)
        
        if response.status_code != 200:
            error_details = f"Status: {response.status_code}, Response: {response.text}"
            print(f"❌ Erreur Azure OpenAI: {error_details}")
            # Retourner les détails de l'erreur au lieu de None
            raise Exception(f"Erreur API Azure OpenAI - {error_details}")
        
        result = response.json()
        markdown_cv = result["choices"][0]["message"]["content"]
        
        print(f"✅ CV généré avec succès ({len(markdown_cv)} caractères)")
        return markdown_cv
        
    except Exception as e:
        print(f"❌ Erreur génération CV: {str(e)}")
        import traceback
        traceback.print_exc()
        raise  # Re-raise l'exception pour que l'appelant puisse la capturer


def convert_markdown_to_word(markdown_cv, email, style, color):
    """
    Convertit un CV Markdown en document Word formaté avec mise en page moderne.
    
    Args:
        markdown_cv: CV au format Markdown
        email: Email du candidat (pour le nom de fichier)
        style: Style du CV
        color: Couleur du CV
    
    Returns:
        str: Chemin du fichier Word généré ou None si erreur
    """
    try:
        # Créer un nouveau document Word
        doc = Document()
        
        # Configuration des marges (1.5cm de chaque côté)
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.6)
            section.bottom_margin = Inches(0.6)
            section.left_margin = Inches(0.8)
            section.right_margin = Inches(0.8)
        
        # Configuration des couleurs par style
        colors_map = {
            'bleu': RGBColor(0, 102, 204),      # Bleu professionnel
            'vert': RGBColor(0, 153, 76),       # Vert moderne
            'gris': RGBColor(52, 73, 94),       # Gris élégant
            'rouge': RGBColor(220, 53, 69)      # Rouge dynamique
        }
        accent_color = colors_map.get(color, colors_map['bleu'])
        light_gray = RGBColor(128, 128, 128)
        
        # Parser le Markdown ligne par ligne
        lines = markdown_cv.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Séparateur horizontal (---)
            if line == '---':
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                run = p.add_run('_' * 80)
                run.font.size = Pt(8)
                run.font.color.rgb = light_gray
                continue
            
            # Titre principal (# ) - NOM
            if line.startswith('# '):
                p = doc.add_paragraph(line[2:].upper())
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.runs[0]
                run.font.size = Pt(28)
                run.font.color.rgb = accent_color
                run.font.name = 'Arial'
                run.bold = True
                p.paragraph_format.space_after = Pt(6)
            
            # Sous-titre (## ) - SECTIONS
            elif line.startswith('## '):
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(6)
                run = p.add_run(line[3:].upper())
                run.font.size = Pt(14)
                run.font.color.rgb = RGBColor(255, 255, 255)  # Blanc
                run.bold = True
                run.font.name = 'Arial'
                
                # Fond coloré pour les titres de section
                from docx.oxml.shared import OxmlElement
                from docx.oxml.ns import qn
                shading_elm = OxmlElement('w:shd')
                # Convertir RGBColor en hexadécimal (ex: RGBColor(0, 112, 192) -> "0070C0")
                hex_color = f"{accent_color[0]:02X}{accent_color[1]:02X}{accent_color[2]:02X}"
                shading_elm.set(qn('w:fill'), hex_color)
                p._element.get_or_add_pPr().append(shading_elm)
            
            # Sous-sous-titre (### ) - Postes/Diplômes
            elif line.startswith('### '):
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(8)
                p.paragraph_format.space_after = Pt(4)
                run = p.add_run(line[4:])
                run.font.size = Pt(12)
                run.font.color.rgb = accent_color
                run.bold = True
                run.font.name = 'Arial'
            
            # Liste à puces (- ou *)
            elif line.startswith('- ') or line.startswith('* '):
                p = doc.add_paragraph(style='List Bullet')
                p.paragraph_format.left_indent = Inches(0.3)
                p.paragraph_format.space_after = Pt(3)
                
                # Parser le texte avec gras et émojis
                text = line[2:]
                if '**' in text:
                    parts = re.split(r'(\*\*.*?\*\*)', text)
                    for part in parts:
                        if part.startswith('**') and part.endswith('**'):
                            run = p.add_run(part[2:-2])
                            run.bold = True
                            run.font.color.rgb = accent_color
                        else:
                            run = p.add_run(part)
                            run.font.size = Pt(10)
                else:
                    run = p.add_run(text)
                    run.font.size = Pt(10)
            
            # Texte avec gras (**texte**) ou émojis
            elif '**' in line or '•' in line:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(3)
                parts = re.split(r'(\*\*.*?\*\*)', line)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                        run.font.color.rgb = accent_color
                        run.font.size = Pt(11)
                    else:
                        run = p.add_run(part)
                        run.font.size = Pt(10)
            
            # Texte normal
            else:
                p = doc.add_paragraph(line)
                p.paragraph_format.space_after = Pt(3)
                run = p.runs[0] if p.runs else None
                if run:
                    run.font.size = Pt(10)
                    run.font.name = 'Arial'
        
        # Sauvegarder le fichier
        filename = f"CV_{email.split('@')[0]}_{uuid.uuid4().hex[:8]}.docx"
        # Utiliser le répertoire temporaire du système (compatible Windows/Linux)
        import tempfile
        temp_dir = tempfile.gettempdir()
        filepath = os.path.join(temp_dir, filename)
        
        doc.save(filepath)
        print(f"✅ CV Word créé: {filepath}")
        
        return filepath
        
    except Exception as e:
        print(f"❌ Erreur conversion Word: {str(e)}")
        import traceback
        traceback.print_exc()
        return None


def send_cv_email(email, word_file_path):
    """
    Envoie le CV par email via Azure Communication Services.
    
    Args:
        email: Adresse email de destination
        word_file_path: Chemin du fichier Word
    
    Returns:
        dict: Résultat de l'envoi
    """
    try:
        if not AZURE_COMMUNICATION_EMAIL_CONNECTION_STRING:
            return {
                "status": "error",
                "message": "Configuration Azure Communication Email manquante"
            }
        
        # Créer le client email
        email_client = EmailClient.from_connection_string(AZURE_COMMUNICATION_EMAIL_CONNECTION_STRING)
        
        # Lire le fichier Word en base64
        with open(word_file_path, 'rb') as f:
            file_content = f.read()
        
        import base64
        file_base64 = base64.b64encode(file_content).decode('utf-8')
        
        # Construire le message
        message = {
            "senderAddress": AZURE_COMMUNICATION_EMAIL_SENDER,
            "recipients": {
                "to": [{"address": email}]
            },
            "content": {
                "subject": "Votre CV professionnel - Waka AI",
                "html": f"""
                <html>
                <body style="font-family: Arial, sans-serif;">
                    <h2 style="color: #0070C0;">Votre CV professionnel</h2>
                    <p>Bonjour,</p>
                    <p>Veuillez trouver ci-joint votre CV professionnel généré par Waka AI.</p>
                    <p>Le CV a été créé à partir des informations que vous avez partagées lors de notre conversation.</p>
                    <p><strong>Conseils pour votre recherche d'emploi:</strong></p>
                    <ul>
                        <li>Personnalisez votre CV pour chaque candidature</li>
                        <li>Mettez en avant vos compétences clés</li>
                        <li>Soyez précis sur vos réalisations</li>
                        <li>Relisez attentivement avant d'envoyer</li>
                    </ul>
                    <p>Bonne chance dans votre recherche d'emploi !</p>
                    <br>
                    <p style="color: #666; font-size: 12px;">Cordialement,<br>L'équipe Waka AI</p>
                </body>
                </html>
                """
            },
            "attachments": [
                {
                    "name": os.path.basename(word_file_path),
                    "contentType": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    "contentInBase64": file_base64
                }
            ]
        }
        
        # Envoyer l'email
        poller = email_client.begin_send(message)
        result = poller.result()
        
        print(f"✅ Email envoyé à {email} - ID: {result['id']}")
        
        return {
            "status": "success",
            "message": f"CV envoyé avec succès à {email}",
            "email_id": result['id']
        }
        
    except Exception as e:
        print(f"❌ Erreur envoi email: {str(e)}")
        import traceback
        traceback.print_exc()
        return {
            "status": "error",
            "message": f"Erreur lors de l'envoi du CV: {str(e)}"
        }


def create_cv(call_id, email, style="moderne", color="bleu"):
    """
    Génère un CV à partir de l'historique de conversation Voice Live.
    
    Args:
        call_id: ID de l'appel Voice Live
        email: Adresse email de destination
        style: Style du CV (classique, moderne, minimaliste)
        color: Couleur du CV (bleu, vert, gris, rouge)
    
    Returns:
        dict: Résultat avec status et détails
    """
    try:
        # 1. Récupérer l'historique depuis Cosmos DB
        print(f"📚 Récupération de l'historique pour call_id: {call_id}")
        conversation_history = get_conversation_history_from_cosmosdb(call_id)
        
        if not conversation_history:
            return {
                "status": "error",
                "message": f"Aucun historique trouvé pour l'appel {call_id}. Assurez-vous d'avoir eu une conversation avant de créer un CV."
            }
        
        print(f"✅ Historique récupéré: {len(conversation_history)} messages")
        
        # 2. Générer le CV
        print("🤖 Génération du CV par IA...")
        try:
            markdown_cv = generate_cv_from_conversation(conversation_history, style, color)
            print(f"✅ CV Markdown généré: {len(markdown_cv)} caractères")
        except Exception as gen_error:
            error_msg = str(gen_error)
            print(f"❌ Erreur lors de la génération: {error_msg}")
            return {
                "status": "error",
                "message": f"Échec de la génération du CV par l'IA: {error_msg}",
                "details": {
                    "error_type": type(gen_error).__name__,
                    "error_message": error_msg,
                    "deployment": AZURE_OPENAI_SUMMARY_DEPLOYMENT,
                    "endpoint": AZURE_OPENAI_SUMMARY_ENDPOINT
                }
            }
        
        # 3. Convertir en Word
        print("📄 Conversion en Word...")
        word_file_path = convert_markdown_to_word(markdown_cv, email, style, color)
        
        if not word_file_path:
            return {
                "status": "partial_success",
                "message": f"CV généré mais échec de la conversion Word. Contenu disponible en Markdown.",
                "call_id": call_id,
                "cv_preview": markdown_cv[:500] + "..." if len(markdown_cv) > 500 else markdown_cv,
                "email": email
            }
        
        # 4. Envoyer par email
        print(f"📧 Envoi par email à {email}...")
        email_result = send_cv_email(email, word_file_path)
        
        # Supprimer le fichier temporaire
        try:
            os.remove(word_file_path)
            print(f"🗑️ Fichier temporaire supprimé: {word_file_path}")
        except:
            pass
        
        if email_result.get('status') == 'success':
            return {
                "status": "success",
                "message": f"✅ CV généré et envoyé avec succès à {email}. Vérifiez votre boîte de réception (et spam).",
                "call_id": call_id,
                "style": style,
                "color": color,
                "email": email,
                "email_id": email_result.get('email_id'),
                "cv_preview": markdown_cv[:300] + "..." if len(markdown_cv) > 300 else markdown_cv
            }
        else:
            return {
                "status": "partial_success",
                "message": f"CV généré mais échec de l'envoi email: {email_result.get('message')}",
                "call_id": call_id,
                "cv_preview": markdown_cv[:500] + "..." if len(markdown_cv) > 500 else markdown_cv,
                "email": email
            }
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        return {
            "status": "error",
            "message": f"Erreur lors de la génération du CV: {str(e)}"
        }


def execute(arguments):
    """Point d'entrée pour l'exécution du tool"""
    call_id = arguments.get("call_id", "")
    email = arguments.get("email", "")
    style = arguments.get("style", "moderne")
    color = arguments.get("color", "bleu")
    
    if not call_id:
        return {
            "status": "error",
            "message": "call_id manquant. Impossible de récupérer l'historique de conversation."
        }
    
    if not email:
        return {
            "status": "error",
            "message": "Adresse email manquante. Veuillez fournir une adresse email pour recevoir le CV."
        }
    
    return create_cv(call_id, email, style, color)
